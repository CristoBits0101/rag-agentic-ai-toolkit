# --- DEPENDENCIAS ---
import hashlib
import pickle
from datetime import datetime
from datetime import timedelta
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.docchat_config import CACHE_DIR
from config.docchat_config import CACHE_EXPIRE_DAYS
from config.docchat_config import MAX_TOTAL_SIZE

try:
    from docling.document_converter import DocumentConverter
except Exception:
    DocumentConverter = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None


class DocumentProcessor:
    def __init__(self):
        self.headers_to_split_on = [
            ("#", "Header 1"),
            ("##", "Header 2"),
            ("###", "Header 3"),
        ]
        self.cache_dir = CACHE_DIR
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def validate_files(self, files: list) -> None:
        total_size = sum(Path(file.name).stat().st_size for file in files)
        if total_size > MAX_TOTAL_SIZE:
            raise ValueError("Uploaded files exceed the maximum total size.")

    def process(self, files: list) -> list[Document]:
        self.validate_files(files)
        all_chunks: list[Document] = []
        seen_hashes: set[str] = set()

        for file in files:
            file_path = Path(file.name)
            file_hash = self._generate_hash(file_path.read_bytes())
            cache_path = self.cache_dir / f"{file_hash}.pkl"

            if self._is_cache_valid(cache_path):
                chunks = self._load_from_cache(cache_path)
            else:
                chunks = self._process_file(file)
                self._save_to_cache(chunks, cache_path)

            for chunk in chunks:
                chunk_hash = self._generate_hash(chunk.page_content.encode("utf-8"))
                if chunk_hash not in seen_hashes:
                    seen_hashes.add(chunk_hash)
                    all_chunks.append(chunk)

        return all_chunks

    def _process_file(self, file) -> list[Document]:
        file_path = Path(file.name)
        suffix = file_path.suffix.lower()
        if suffix not in {".pdf", ".docx", ".txt", ".md"}:
            return []

        markdown_text = self._extract_markdown(file_path)
        return self._split_markdown(markdown_text, source=file_path.name)

    def _extract_markdown(self, file_path: Path) -> str:
        if DocumentConverter is not None:
            try:
                converter = DocumentConverter()
                result = converter.convert(str(file_path))
                document = getattr(result, "document", result)
                if hasattr(document, "export_to_markdown"):
                    return document.export_to_markdown()
            except Exception:
                pass

        suffix = file_path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".pdf" and PdfReader is not None:
            reader = PdfReader(str(file_path))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".docx" and DocxDocument is not None:
            document = DocxDocument(str(file_path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _split_markdown(self, markdown_text: str, source: str) -> list[Document]:
        markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=self.headers_to_split_on)
        split_documents = markdown_splitter.split_text(markdown_text)

        if not split_documents:
            splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=120)
            return splitter.create_documents([markdown_text], metadatas=[{"source": source}])

        documents: list[Document] = []
        fallback_splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=120)
        for split_document in split_documents:
            metadata = {"source": source, **split_document.metadata}
            if len(split_document.page_content) > 1200:
                documents.extend(
                    fallback_splitter.create_documents([split_document.page_content], metadatas=[metadata])
                )
            else:
                documents.append(Document(page_content=split_document.page_content, metadata=metadata))
        return documents

    def _generate_hash(self, content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()

    def _save_to_cache(self, chunks: list[Document], cache_path: Path) -> None:
        with cache_path.open("wb") as handle:
            pickle.dump({"timestamp": datetime.utcnow(), "chunks": chunks}, handle)

    def _load_from_cache(self, cache_path: Path) -> list[Document]:
        with cache_path.open("rb") as handle:
            payload = pickle.load(handle)
        return payload["chunks"]

    def _is_cache_valid(self, cache_path: Path) -> bool:
        if not cache_path.exists():
            return False
        modified_time = datetime.utcfromtimestamp(cache_path.stat().st_mtime)
        return modified_time > datetime.utcnow() - timedelta(days=CACHE_EXPIRE_DAYS)